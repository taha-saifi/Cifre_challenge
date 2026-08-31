#!/usr/bin/env python3
"""Render deliverables/note.md to .docx and MEASURE its real page count.

The 4-page limit of §18.1 is a hard submission constraint, and a word-count estimate is
not a measurement: pagination depends on font, margins, tables and widow control. This
script therefore builds the document and then asks Word itself how many pages it has,
via COM. If Word is unavailable the script says so rather than substituting an estimate.

The AI Usage & Verification Log starts on its own page, because the brief counts the
4 pages "hors références et hors AI Usage Log" -- so the body page count is the count
before that break, and it is reported separately.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Pt, RGBColor, Cm

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "deliverables" / "note.md"
OUT = ROOT / "deliverables" / "note.docx"

ACCENT = RGBColor(0x0E, 0x74, 0x90)
MUTED = RGBColor(0x5A, 0x64, 0x72)
LOG_HEADING = "AI Usage & Verification Log"


def add_runs(paragraph, text: str) -> None:
    """Render inline **bold**, *italic* and `code` into a paragraph."""
    for token in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)", text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            paragraph.add_run(token[1:-1]).italic = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
        else:
            paragraph.add_run(token)


def build() -> Document:
    doc = Document()

    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21), Cm(29.7)   # A4
    for side in ("left_margin", "right_margin"):
        setattr(section, side, Cm(2.0))
    section.top_margin, section.bottom_margin = Cm(1.8), Cm(1.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    lines = SRC.read_text(encoding="utf-8").split("\n")
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()

        if line.startswith("| ") and index + 1 < len(lines) and set(lines[index + 1]) <= set("|-: "):
            rows = []
            while index < len(lines) and lines[index].startswith("|"):
                cells = [c.strip() for c in lines[index].strip().strip("|").split("|")]
                if not set("".join(cells)) <= set("-: "):
                    rows.append(cells)
                index += 1
            table = doc.add_table(rows=0, cols=len(rows[0]))
            table.style = "Light Grid Accent 1"
            for r, row in enumerate(rows):
                cells = table.add_row().cells
                for c, value in enumerate(row[:len(cells)]):
                    cells[c].text = ""
                    para = cells[c].paragraphs[0]
                    para.paragraph_format.space_after = Pt(1)
                    add_runs(para, value)
                    for run in para.runs:
                        run.font.size = Pt(8)
                        if r == 0:
                            run.bold = True
            doc.add_paragraph()
            continue

        if line.startswith("# "):
            para = doc.add_paragraph()
            run = para.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(15)
            run.font.color.rgb = ACCENT
            para.paragraph_format.space_after = Pt(6)
        elif line.startswith("## "):
            title = line[3:]
            if LOG_HEADING in title:
                # The 4-page limit excludes this section: start it on a fresh page so the
                # body count is unambiguous.
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            para = doc.add_paragraph()
            run = para.add_run(title)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = ACCENT
            para.paragraph_format.space_before = Pt(7)
            para.paragraph_format.space_after = Pt(3)
        elif line.startswith("**Note de proposition"):
            para = doc.add_paragraph()
            add_runs(para, line)
            for run in para.runs:
                run.font.color.rgb = MUTED
                run.font.size = Pt(9)
        elif re.match(r"^[-*] ", line):
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.space_after = Pt(2)
            add_runs(para, line[2:])
        elif re.match(r"^\d+\. ", line):
            para = doc.add_paragraph(style="List Number")
            para.paragraph_format.space_after = Pt(2)
            add_runs(para, re.sub(r"^\d+\.\s*", "", line))
        elif line.startswith("---") or not line.strip():
            pass
        else:
            buffer = [line]
            while index + 1 < len(lines) and lines[index + 1].strip() and not re.match(
                    r"^(#|\||[-*] |\d+\. |---)", lines[index + 1]):
                index += 1
                buffer.append(lines[index].strip())
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs(para, " ".join(buffer))
        index += 1

    return doc


def measure_pages(path: Path) -> tuple[int | None, int | None]:
    """Ask Word for the real page count, and for the page the AI log starts on.

    Returns (total_pages, body_pages) or (None, None) when Word is unavailable.
    """
    try:
        import win32com.client as com
    except ImportError:
        return None, None
    word = None
    try:
        word = com.Dispatch("Word.Application")
        word.Visible = False
        document = word.Documents.Open(str(path), ReadOnly=True)
        document.Repaginate()
        total = int(document.ComputeStatistics(2))  # wdStatisticPages
        body = total
        for paragraph in document.Paragraphs:
            if LOG_HEADING in paragraph.Range.Text:
                body = int(paragraph.Range.Information(3)) - 1  # wdActiveEndPageNumber
                break
        document.Close(False)
        return total, body
    except Exception as exc:  # noqa: BLE001
        print(f"  (Word indisponible : {type(exc).__name__}: {exc})")
        return None, None
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:  # noqa: BLE001
                pass


def main() -> int:
    build().save(OUT)
    print(f"wrote {OUT}")
    total, body = measure_pages(OUT)
    if total is None:
        print("PAGES : non mesurables ici (Word absent). Ne pas déclarer un nombre de pages.")
        return 1
    print(f"PAGES mesurees par Word : {total} au total, {body} pour le corps "
          f"(hors AI Usage Log)")
    print("Contrainte §18.1 : corps <= 4 pages -> "
          f"{'CONFORME' if body <= 4 else 'DEPASSEMENT de ' + str(body - 4) + ' page(s)'}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
